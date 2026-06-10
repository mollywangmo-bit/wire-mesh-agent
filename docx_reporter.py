"""
丝网行业研究 Agent - Word (.docx) 报告生成器

直接从 Markdown 渲染为 Word 文档，保留：
- 真表格（列对齐、表头着色）
- 超链接（可点击）
- 相关性标签着色
- 图表附录（着色表格模拟柱状条）

纯 Python（python-docx），零系统库依赖。
"""
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── 相关性标签配色 ──────────────────────────────────────────────
TAG_EMOJIS = {"🔗", "⚡", "📡"}
TAG_COLORS: dict[str, RGBColor] = {
    "🔗": RGBColor(0x00, 0x70, 0xF3),  # 蓝色 — 直接相关
    "⚡": RGBColor(0x46, 0xA7, 0x58),  # 绿色 — 间接影响
    "📡": RGBColor(0xFF, 0xB2, 0x24),  # 橙色 — 趋势信号
}
TAG_LABELS: dict[str, str] = {
    "🔗": "直接相关",
    "⚡": "间接影响",
    "📡": "趋势信号",
}
_SVG_COLORS = [
    "#0070F3", "#46A758", "#FFB224", "#E5484D", "#8B5CF6",
    "#06B6D4", "#F472B6", "#84CC16", "#F97316", "#6366F1",
    "#14B8A6",
]


# ── 辅助函数 ────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str):
    """设置单元格底色"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold: bool = False,
                   color: Optional[RGBColor] = None, size: int = 10):
    """设置单元格纯文本"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_hyperlink(paragraph, text: str, url: str):
    """在段落中添加可点击超链接"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:color w:val="0563C1"/>'
        f'      <w:u w:val="single"/>'
        f'      <w:rStyle w:val="Hyperlink"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{_xml_escape(text)}</w:t>'
        f'  </w:r>'
        f'</w:hyperlink>'
    )
    paragraph._p.append(hyperlink)


def _xml_escape(text: str) -> str:
    """XML 转义（python-docx 的 <w:t> 内容需要）"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


# ── Markdown 行内解析 ───────────────────────────────────────────

def _add_formatted_paragraph(doc, text: str, style: Optional[str] = None,
                             indent: float = 0):
    """解析内联标记（粗体、链接、标签）并添加格式化段落"""
    if not text.strip():
        return

    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    if indent:
        p.paragraph_format.left_indent = Cm(indent)

    _populate_paragraph_runs(p, text)
    return p


def _populate_paragraph_runs(p, text: str):
    """将文本按内联标记切分后添加格式化的 Run"""
    # 先清理掉紧跟在 emoji 后的中文标签文本（避免重复）
    for emoji in sorted(TAG_EMOJIS, key=len, reverse=True):
        label = TAG_LABELS.get(emoji, "")
        # "🔗 直接相关" 或 "🔗直接相关"
        text = re.sub(re.escape(emoji) + r"\s*" + re.escape(label), emoji, text)

    pattern = r'(\*\*(.+?)\*\*|\[([^\]]+)\]\(([^)]+)\)|[' + ''.join(TAG_EMOJIS) + r']|[一-鿿\w\d\s.,;:!?\/()（）\-+%#@$&°\[\]]+|.)'
    tokens = re.findall(pattern, text, re.MULTILINE)

    for match in tokens:
        full = match[0]
        # **bold**
        if full.startswith("**") and match[2]:
            run = p.add_run(match[2])
            run.bold = True
        # [text](url)
        elif full.startswith("[") and match[3]:
            _add_hyperlink(p, match[2], match[3])
        # 相关性标签 emoji
        elif full in TAG_EMOJIS:
            run = p.add_run(f" {TAG_LABELS.get(full, full)} ")
            run.font.color.rgb = TAG_COLORS.get(full, RGBColor(0, 0, 0))
            run.bold = True
            run.font.size = Pt(9)
        else:
            if full.strip():
                run = p.add_run(full)


# ── 报告解析 ────────────────────────────────────────────────────

def _extract_title_and_date(report_text: str) -> tuple[str, str]:
    """从首行提取标题和日期"""
    lines = report_text.strip().split("\n")
    title = "丝网行业报告"
    date_str = datetime.now().strftime("%Y-%m-%d")

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title_text = stripped[2:]
            date_match = re.search(r'[｜|]\s*(\d{4}-\d{2}-\d{2})', title_text)
            if date_match:
                date_str = date_match.group(1)
                title_text = title_text.replace(date_match.group(0), "").strip(" ｜|").strip()
            title = title_text
            break

    return title, date_str


# ── 章节渲染 ────────────────────────────────────────────────────

def _render_chapter(doc, title: str, lines: list[str]):
    """渲染一个 ## 章节"""
    doc.add_heading(title, level=2)

    # 收集连续表格行
    table_lines = []

    def _flush_table():
        if len(table_lines) < 2:
            return  # 至少需要表头 + 一行数据
        _add_table(doc, table_lines)
        table_lines.clear()

    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped:
            _flush_table()
            doc.add_paragraph()  # 空行
            continue

        # 表格行
        if stripped.startswith("|") and stripped.endswith("|"):
            if re.match(r"^\|[-:| ]+\|$", stripped):
                continue  # 跳过分隔行
            table_lines.append(stripped)
            continue
        else:
            _flush_table()

        # 分隔线
        if stripped.startswith("---"):
            doc.add_paragraph("―" * 40)
            continue

        # 标题
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            continue

        # 引用
        if stripped.startswith("> "):
            _add_formatted_paragraph(doc, stripped[2:], indent=1.0)
            continue

        # 无序列表
        if stripped.startswith("- "):
            p = _add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
            continue

        # 有序列表
        if re.match(r"^\d+\. ", stripped):
            p = _add_formatted_paragraph(doc, re.sub(r"^\d+\.\s*", "", stripped),
                                         style="List Number")
            continue

        # 普通段落
        _add_formatted_paragraph(doc, stripped)

    _flush_table()


def _add_table(doc, table_lines: list[str]):
    """将 markdown 表格行渲染为 Word 表格"""
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_cells in enumerate(rows_data):
        for ci in range(num_cols):
            cell = table.cell(ri, ci)
            text = row_cells[ci] if ci < len(row_cells) else ""
            # 首行作表头
            if ri == 0:
                _set_cell_text(cell, text, bold=True, size=9)
                _set_cell_shading(cell, "E8E8E8")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                _set_cell_text(cell, text, size=9)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ── 图表附录 ────────────────────────────────────────────────────

def _add_charts_section(doc, sections: list[dict],
                        relevance_counts: Optional[dict[str, int]] = None):
    """用着色表格模拟柱状图"""
    if relevance_counts is None:
        relevance_counts = {"direct": 0, "indirect": 0, "signal": 0}
    doc.add_paragraph()  # 间距
    doc.add_heading("📊 报告概览", level=2)

    # ── 各板块内容量 ──
    doc.add_heading("各板块内容量", level=3)
    if not sections:
        return

    max_words = max(s.get("word_count", 1) for s in sections)
    bar_colors = _SVG_COLORS

    table = doc.add_table(rows=len(sections) + 1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["板块", "内容量", "字数"]
    for ci, h in enumerate(headers):
        _set_cell_text(table.cell(0, ci), h, bold=True, size=9)
        _set_cell_shading(table.cell(0, ci), "F0F0F0")

    for i, sec in enumerate(sections):
        title = sec.get("title", "")
        # 去除编号前缀显示
        short = re.sub(r"^\d+\.\s*", "", title)[:18]
        count = sec.get("word_count", 0)
        bar_width = max(count / max_words, 0.05) if max_words else 0.05
        color = bar_colors[i % len(bar_colors)]

        _set_cell_text(table.cell(i + 1, 0), short, size=9)
        _set_cell_text(table.cell(i + 1, 2), str(count), size=9)

        # 内容量列：用着色单元格模拟柱状条
        bar_cell = table.cell(i + 1, 1)
        _set_cell_shading(bar_cell, color)
        # 如果柱状条够宽，显示百分比文字
        if bar_width > 0.15:
            _set_cell_text(bar_cell, f"{int(bar_width * 100)}%", color=RGBColor(0xFF, 0xFF, 0xFF), size=8)
        else:
            _set_cell_text(bar_cell, "", size=8)

        # 设置列宽：板块 / 柱状条 / 字数
        for row in table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(8)
            row.cells[2].width = Cm(2)

    # ── 信息关联度分布 ──
    doc.add_paragraph()
    doc.add_heading("信息关联度分布", level=3)

    rtable = doc.add_table(rows=4, cols=3)
    rtable.style = "Table Grid"
    rtable.alignment = WD_TABLE_ALIGNMENT.CENTER

    rheaders = ["类型", "数量", "占比"]
    for ci, h in enumerate(rheaders):
        _set_cell_text(rtable.cell(0, ci), h, bold=True, size=9)
        _set_cell_shading(rtable.cell(0, ci), "F0F0F0")

    categories = [
        ("direct", "直接相关", "🔗"),
        ("indirect", "间接影响", "⚡"),
        ("signal", "趋势信号", "📡"),
    ]
    total = sum(relevance_counts.values()) or 1

    for i, (key, label, emoji) in enumerate(categories):
        count = relevance_counts.get(key, 0)
        pct = f"{count / total * 100:.0f}%"

        color = TAG_COLORS.get(emoji, RGBColor(0, 0, 0))
        _set_cell_text(rtable.cell(i + 1, 0), f"{emoji} {label}", size=9)
        _set_cell_text(rtable.cell(i + 1, 1), str(count), size=9)
        _set_cell_text(rtable.cell(i + 1, 2), pct, size=9)
        # 用淡色底色
        hex_color = str(color)
        _set_cell_shading(rtable.cell(i + 1, 0), f"{hex_color}20")

    for row in rtable.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(2)
        row.cells[2].width = Cm(2)


# ── 主入口 ──────────────────────────────────────────────────────

def generate_docx(report_text: str, output_path: str | Path) -> str:
    """从 Markdown 报告生成 Word (.docx) 文件

    Args:
        report_text: 报告全文（Markdown 格式）
        output_path: 输出 .docx 文件路径

    Returns:
        生成的 .docx 文件路径
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── 设置默认中文字体 ──
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    font = style.font
    font.name = "等线"
    # 设置东亚字体
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:eastAsia"), "等线")

    # ── 解析标题 ──
    title, date_str = _extract_title_and_date(report_text)

    # ── 文档头部 ──
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run(f"丝网行业报告｜{date_str}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xF3)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run(title)
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # 间距

    # ── 按 ## 拆分段落 ──
    lines = report_text.split("\n")
    sections_data = []  # 存 (title, body_lines, word_count)
    current_title = None
    current_lines = []
    relevance_counts: dict[str, int] = {"direct": 0, "indirect": 0, "signal": 0}

    # 扫描 emoji 计数
    tag_to_key = {"🔗": "direct", "⚡": "indirect", "📡": "signal"}
    def _count_tags(text: str):
        for emoji, key in tag_to_key.items():
            if emoji in text:
                relevance_counts[key] += text.count(emoji)

    for line in lines:
        stripped = line.strip()
        _count_tags(stripped)

        if stripped.startswith("## ") and not stripped.startswith("### "):
            if current_title:
                wc = sum(len(l.strip()) for l in current_lines if l.strip())
                sections_data.append((current_title, list(current_lines), wc))
            current_title = stripped[3:]
            current_lines = []
        elif stripped.startswith("# ") and not stripped.startswith("## "):
            continue  # 主标题跳过
        else:
            if current_title is not None:
                current_lines.append(line)
            # 主标题前的元信息行（日期等）忽略

    if current_title:
        wc = sum(len(l.strip()) for l in current_lines if l.strip())
        sections_data.append((current_title, list(current_lines), wc))

    # ── 渲染每个段落 ──
    for sec_title, sec_lines, _wc in sections_data:
        _render_chapter(doc, sec_title, sec_lines)

    # ── 图表附录 ──
    chart_sections = [
        {"title": t, "word_count": wc}
        for t, _lines, wc in sections_data
    ]
    _add_charts_section(doc, chart_sections, relevance_counts)

    # ── 页脚信息 ──
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"—— 由丝网行业研究 Agent 自动生成 ｜ {date_str} ——"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── 保存 ──
    doc.save(str(output_path))
    return str(output_path)


if __name__ == "__main__":
    test = """# 丝网行业周报｜2026-06-05

## 1. 本周一句话判断
本周不锈钢价格震荡上行，[青山集团](https://example.com)上调304冷轧出厂价200元/吨。

## 2. 原材料价格

| 品种 | 价格 | 涨跌 |
|------|------|------|
| 304不锈钢 | 15200元/吨 | +1.3% |
| LME镍 | 19200美元/吨 | -0.7% |

- **青山集团**304冷轧卷板 🔗 直接相关 — 连续三周上涨
- LME镍收盘19200美元/吨
- 碳纤维复合材料 ⚡ 间接影响

## 3. 值得关注的新闻

1. 青山集团印尼镍冶炼产能扩建获批 🔗 直接相关
2. 日本精線精密丝材业务增长12%
"""
    path = generate_docx(test, "/tmp/test_report.docx")
    print(f"Word 已生成: {path}")
    print(f"大小: {__import__('os').path.getsize(path) / 1024:.0f} KB")
